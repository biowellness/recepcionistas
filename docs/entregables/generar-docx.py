#!/usr/bin/env python3
"""
Genera los entregables .docx del bloque de gestión de sesiones:
  - BioWellness-Gestion-de-Sesiones-Andres.docx  (resumen para Andrés)
  - BioWellness-Guia-Recepcion.docx              (guía para recepcionistas)

Uso: python3 docs/entregables/generar-docx.py
Requiere: python-docx  (pip install python-docx)
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

TEAL = RGBColor(0x0C, 0x8A, 0x6E)
GRIS = RGBColor(0x66, 0x66, 0x66)
OUT = os.path.dirname(os.path.abspath(__file__))


def base(doc: Document):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)


def titulo(doc, texto, sub):
    h = doc.add_paragraph()
    r = h.add_run(texto)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = TEAL
    p = doc.add_paragraph()
    rs = p.add_run(sub)
    rs.font.size = Pt(11)
    rs.font.color.rgb = GRIS
    doc.add_paragraph()


def h1(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = TEAL
    return p


def h2(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.bold = True
    r.font.size = Pt(12)
    return p


def parrafo(doc, texto):
    return doc.add_paragraph(texto)


def bullet(doc, texto, neg=None):
    p = doc.add_paragraph(style="List Bullet")
    if neg:
        r = p.add_run(neg)
        r.bold = True
        p.add_run(texto)
    else:
        p.add_run(texto)
    return p


def nota(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run("Nota: ")
    r.bold = True
    r.font.color.rgb = TEAL
    p.add_run(texto)
    return p


# ───────────────────────── Doc para Andrés ─────────────────────────
def doc_andres():
    doc = Document()
    base(doc)
    titulo(
        doc,
        "BioWellness · Gestión de Sesiones",
        "Resumen del bloque entregado · Recepción (Bloque 0) · 21/06/2026",
    )

    h1(doc, "En una frase")
    parrafo(
        doc,
        "Cerramos el circuito para que ninguna sesión paga se pierda: el sistema "
        "muestra el saldo de cada cliente, ayuda a dejar agendado el mes de una vez "
        "y avisa solo (WhatsApp + email) antes de cada turno y cuando hay sesiones "
        "por vencer.",
    )

    h1(doc, "Qué se entregó")

    h2(doc, "1. Tablero “Planes y sesiones”")
    bullet(doc, "Lista a todos los clientes con plan activo y reparte sus sesiones en tres grupos: realizadas, agendadas a futuro y libres (sin agendar).")
    bullet(doc, "Ordena por urgencia (membresía: cierre de mes; paquete: vencimiento) y permite filtrar “En riesgo”.")
    bullet(doc, "Desde cada fila se pasa directo a atender al cliente.")

    h2(doc, "2. Pre-agenda de membresías")
    bullet(doc, "Al asignar una membresía, propone y reserva de una vez las sesiones del mes según la frecuencia (2 o 3 por semana).")
    bullet(doc, "Recepción elige días y horario; el sistema asigna sala, valida las reglas y descuenta del plan. Manda un único WhatsApp de resumen.")

    h2(doc, "3. Recordatorios automáticos")
    bullet(doc, "Aviso 24 h y 1 h antes de cada turno confirmado (los combos cuentan como una sola visita).")
    bullet(doc, "Aviso de “saldo en riesgo” cuando quedan sesiones por perderse pronto.")
    bullet(doc, "Por WhatsApp y email. No se repiten: cada aviso se manda una sola vez.")

    h2(doc, "4. Modo oscuro / claro")
    bullet(doc, "Botón en la pantalla para alternar tema; queda recordado en cada equipo.")

    h1(doc, "Calidad y confiabilidad")
    bullet(doc, "114 pruebas automáticas en verde; la lógica de negocio (saldos, fechas, reglas) está aislada y testeada.")
    bullet(doc, "Toda la lógica vive en el backend: la recepción no calcula ni decide nada que el sistema pueda calcular.")
    bullet(doc, "Los avisos son idempotentes (no spamean) y quedan registrados aunque el canal esté caído.")

    h1(doc, "Qué falta para encenderlo (acción del equipo)")
    parrafo(
        doc,
        "El software está listo y desplegado. Hoy los mensajes se registran pero "
        "todavía no se envían, hasta completar las cuentas:",
    )
    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = "Light Grid Accent 1"
    hdr = tabla.rows[0].cells
    hdr[0].paragraphs[0].add_run("Pendiente").bold = True
    hdr[1].paragraphs[0].add_run("Para qué").bold = True
    filas = [
        ("Cuenta WhatsApp Business (Twilio) + plantilla aprobada", "Enviar confirmaciones y recordatorios por WhatsApp."),
        ("Remitente verificado en AWS SES", "Enviar los recordatorios por email."),
        ("Activar el horario del bot de recordatorios", "Que el aviso salga solo cada hora (configuración de una sola vez)."),
    ]
    for a, b in filas:
        c = tabla.add_row().cells
        c[0].text = a
        c[1].text = b
    doc.add_paragraph()

    h1(doc, "Decisiones de negocio aún abiertas")
    bullet(doc, "Tabla de contraindicaciones: hoy hay un borrador estándar; requiere validación del Director Médico.", neg="Clínico — ")
    bullet(doc, "Precio de consulta del Dr. Conrado: provisorio (ARS 150.000) a confirmar.", neg="Catálogo — ")
    bullet(doc, "Detalles de IHHT (Express/Premium) y algunos descuentos v9 a confirmar con el equipo.", neg="Catálogo — ")

    nota(
        doc,
        "El detalle técnico está en el repositorio: docs/app-recepcion.md (la app), "
        "docs/bots.md (automatizaciones y secretos) y docs/decisiones-pendientes.md.",
    )
    ruta = os.path.join(OUT, "BioWellness-Gestion-de-Sesiones-Andres.docx")
    doc.save(ruta)
    return ruta


# ─────────────────── Guía para recepcionistas ───────────────────
def doc_recepcion():
    doc = Document()
    base(doc)
    titulo(
        doc,
        "Guía rápida de Recepción",
        "Planes, sesiones y recordatorios · BioWellness",
    )

    parrafo(
        doc,
        "Esta guía explica las funciones nuevas. Regla de oro: vos no calculás "
        "precios ni saldos — el sistema lo hace. Tu trabajo es atender, agendar "
        "y acompañar al cliente.",
    )

    h1(doc, "Las pestañas de arriba")
    bullet(doc, "el día de hoy por sala; tocás un hueco libre para reservar.", neg="Agenda: ")
    bullet(doc, "todos los clientes con plan y cuántas sesiones les quedan.", neg="Planes y sesiones: ")
    bullet(doc, "buscás a una persona y ves su ficha (plan, reservas, cobro).", neg="Atender paciente: ")
    bullet(doc, "números de gestión.", neg="Reportes: ")

    h1(doc, "Planes y sesiones (lo más importante)")
    parrafo(doc, "Cada fila es un cliente con plan. La barra de colores muestra sus sesiones:")
    bullet(doc, "ya las hizo.", neg="Realizadas: ")
    bullet(doc, "tiene turno reservado a futuro.", neg="Agendadas: ")
    bullet(doc, "le quedan sin reservar (¡estas son las que hay que agendar!).", neg="Libres: ")
    parrafo(doc, "El cartelito de color a la derecha avisa la urgencia:")
    bullet(doc, "rojo/naranja = quedan sesiones y se vencen pronto (cierre de mes o vencimiento del paquete).")
    bullet(doc, "Usá el filtro “En riesgo” para ver solo los urgentes y llamarlos.")
    bullet(doc, "El botón “Atender” te lleva directo a la ficha de esa persona para agendar.")
    nota(doc, "Las sesiones no usadas se pierden: la membresía se reinicia cada mes y el paquete vence. Por eso conviene agendarlas a tiempo.")

    h1(doc, "Pre-agendar el mes de una membresía")
    parrafo(doc, "Cuando asignás (o abrís) una membresía con saldo, aparece el botón “Pre-agendar mes”. Sirve para dejar todo el mes reservado de una sola vez:")
    bullet(doc, "Tocá “Pre-agendar mes”.")
    bullet(doc, "Elegí los días de la semana (ya vienen sugeridos según la frecuencia) y el horario.")
    bullet(doc, "Mirá la lista de fechas propuestas y tocá “Reservar”.")
    bullet(doc, "El sistema reserva una por una y te muestra cuáles entraron (✓) y cuáles no (✗, por ej. sin sala). Las que fallen las agendás a mano.")
    nota(doc, "El cliente recibe un solo mensaje de resumen, no uno por cada sesión.")

    h1(doc, "Recordatorios automáticos")
    parrafo(doc, "Esto lo hace el sistema solo, no tenés que hacer nada:")
    bullet(doc, "Avisa al cliente 24 horas y 1 hora antes de su turno.")
    bullet(doc, "Avisa cuando le quedan sesiones por vencer, invitando a agendar.")
    bullet(doc, "Los manda por WhatsApp y email, y nunca repite el mismo aviso.")

    h1(doc, "Modo oscuro / claro")
    parrafo(doc, "Arriba a la derecha, al lado de “Salir”, hay un botón con un sol o una luna. Tocalo para cambiar entre pantalla clara y oscura. Queda guardado para la próxima vez.")

    h1(doc, "Si algo no cierra")
    bullet(doc, "Un turno con plan no descuenta del saldo o el saldo se ve raro: avisá al equipo técnico (no lo fuerces a mano).")
    bullet(doc, "El sistema bloquea una reserva: te dice el motivo (sala ocupada, sin saldo, etc.). Seguí ese mensaje.")

    ruta = os.path.join(OUT, "BioWellness-Guia-Recepcion.docx")
    doc.save(ruta)
    return ruta


if __name__ == "__main__":
    print("Generado:", doc_andres())
    print("Generado:", doc_recepcion())
